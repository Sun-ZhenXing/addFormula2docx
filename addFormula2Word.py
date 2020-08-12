from docx import Document
from lxml import etree
import latex2mathml.converter

def getTransform(filepath):
    xslt = etree.parse(filepath)
    return etree.XSLT(xslt)

def addEq2Word(transform, formula, eq_cls, paragraph, math_format="in-line"):
    """Add Math Formula(latex/mathml) to docx
    ---------------------------------------
    | parameter|      content         |
    |----------|----------------------|
    | transform|(func) etree.XSLT     |
    | formula  |(str)string of formula|
    | eq_cls   | (str) type of formula["latex", "mathml"]|
    | paragraph| (obj) docx.paragraph|
    | [math_format]| (str) in line or not["in-line", "non-in-line"]|
    
    >>> from docx import Document
    >>> doc = Document()
    >>> para1 = doc.add_paragraph("Paragraph1, formula:")
    >>> transform = getTransform("MML2OMML.XSL")
    >>> addEq2Word(transform, MathML_Str, "mathml", para1)
    >>> addEq2Word(transform, Latex_Str, "mathml", doc.add_paragraph(), "non-in-line")
    >>> doc.save("test.docx")
    """
    if eq_cls == "latex":
        mathml = latex2mathml.converter.convert(formula)
    elif eq_cls == "mathml":
        mathml = formula
    else:
        raise ValueError(f"not support formula type: {eq_cls}")
    tree = etree.fromstring(mathml)
    new_dom = transform(tree)
    if math_format == "in-line":
        paragraph._element.append(new_dom.getroot())
    elif math_format == "non-in-line":
        wrapper = etree.fromstring('<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '</m:oMathPara>')
        wrapper.append(new_dom.getroot())
        paragraph._element.append(wrapper)
    else:
        raise ValueError(f"Unknown format: {math_format}")

if __name__ == "__main__":
    mathml_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"><mfrac><mn>1</mn><mn>2</mn></mfrac></math>'
    doc = Document()
    para1 = doc.add_paragraph("Text ... Text")
    transform = getTransform("MML2OMML.XSL")
    para2 = doc.add_paragraph()
    addEq2Word(transform, mathml_string, "mathml", para1)
    addEq2Word(transform, mathml_string, "mathml", para2, "non-in-line")
    a=addEq2Word(transform, mathml_string, "mathml", doc.add_paragraph(), "non-in-line")
    doc.save("demo.docx")
