import docx.document
import docx.text.paragraph
import re
from lxml import etree

LaTeX = 0
MathML = 1
OMML = 2

Transforms = {}

# Transforms = {
#     (LaTeX,   MathML) :   func1,
#     (LaTeX,   OMML)   :   func2,
#     (MathML,  OMML)   :   func3,
#     (MathML,  LaTeX)  :   func4,
#     (OMML,    MathML) :   func5,
#     (OMML,    LaTeX)  :   func6,
# }


def init(func='all', **kwargs):
    """
    Import or modify conversion functions

    Usage
    ==============
    Import All available functions:
        >>> from formulas import init, Formula
        >>> init()
    Import specific function:
        >>> init(func, From='latex', To='mathml')
    or:
        >>> from formulas import LaTeX, MathML, OMML
        >>> init(func, From=LaTeX, To=MathML)
    and the 'func' should be defined like this:
        >>> def func(str|etree, display='block'|'inline'):
        >>>     ...
    """
    if func == 'all':
        from .utils import latex_to_mathml, mathml_to_omml, omml_to_mathml
        Transforms[(LaTeX, MathML)] = latex_to_mathml
        Transforms[(MathML, OMML)] = mathml_to_omml
        Transforms[(OMML, MathML)] = omml_to_mathml
    if not(('From' in kwargs) or ('To' in kwargs)):
        raise ValueError("The key 'From', 'To' is required")
    From = kwargs['From']
    To = kwargs['To']
    Transforms[(From, To)] = func
    return Transforms


class Formula:
    """
    Formula Object

    To insert formulas into Word documents(.docx|.doc)

    Usage
    =============
        >>> from formulas import Formulas, init
        >>> from docx import Document
        >>> init()
        >>> doc = Document()
        >>> para = doc.add_paragraph('some text...')
        >>> fm1 = Formula('LaTeX|MathML|OMML string')
        >>> fm2 = Formula('LaTeX string', 'latex')
        >>> fm3 = Formula('MathML string', 'mathml', 'block')
        >>> fm1.add_to(doc)
        >>> para2 = doc.add_paragraph('some text...') + fm2
        >>> para3 = doc + fm3
    other:
        >>> fm1.get_latex(True)
        'LaTeX string if its available...'
        >>> fm1.get_mathml(to_string=True)
        >>> fm1.get_omml(to_string=True, safe_mode=True)
    safe mode not support a formula be converted many times.
    """

    def __init__(self, formula, fm_type='auto', fm_format='inline'):
        self.formula = [None, None, None]
        if fm_type == 'auto':
            self.fm_type = self._getType(formula)
        elif str(fm_type).lower() == 'latex' or fm_type == LaTeX:
            self.fm_type = LaTeX
            self.formula[LaTeX] = formula
        elif str(fm_type).lower() == 'mathml' or fm_type == MathML:
            self.fm_type = MathML
            self.formula[MathML] = etree.fromstring(formula)
        elif str(fm_type).lower() == 'omml' or fm_type == OMML:
            self.fm_type = OMML
            self.formula[OMML] = etree.fromstring(formula)
        else:
            raise ValueError(f"Unknown type '{fm_type}' ")
        if fm_format != 'inline':
            self.fm_format = 'block'
        else:
            self.fm_format = fm_format

    def __add__(self, value):
        if not self.formula[OMML]:
            self.get_omml()
        if not self.get_omml():
            raise ValueError("Can't get the formula of 'OMML' format")
        if type(value) == docx.text.paragraph.Paragraph:
            value._element.append(self.formula[OMML].getroot())
        elif type(value) == docx.document.Document:
            value.add_paragraph()._element.append(self.formula[OMML].getroot())
        else:
            raise ValueError(
                "Only 'Document' & 'Paragraph' object are supported")
        return value

    def __mul__(self, value):
        return self.__add__(value)

    def __rmul__(self, value):
        return self.__add__(value)

    def __radd__(self, value):
        return self.__add__(value)

    def _getType(self, fm_string):
        if re.search(r'https*:\/\/schemas\.openxmlformats\.org\/officeDocument\/',
                     fm_string, re.M):
            return OMML
        if re.search(r'https*:\/\/www\.w3\.org\/1998\/Math\/MathML',
                     fm_string, re.M):
            return MathML
        return LaTeX

    def add_to(self, value):
        """
        formula.add_to(document|paragraph)
        """
        return self.__add__(value)

    def get_formula(self):
        output = []
        if self.formula[LaTeX]:
            output.append('LaTeX')
        if self.formula[MathML]:
            output.append('MathML')
        if self.formula[OMML]:
            output.append('OMML')
        return output

    def get_latex(self, to_string=False, safe_mode=False):
        if self.fm_type == LaTeX:
            if to_string:
                return self.formula[LaTeX]
            else:
                return None
        if (self.fm_type, LaTeX) in Transforms:
            self.formula[LaTeX] = Transforms[(self.fm_type, LaTeX)](
                self.formula[self.fm_type], display=self.fm_format
            )
            if to_string:
                return self.formula[LaTeX]
            else:
                return None
        if safe_mode:
            raise ValueError("Can't get the formula of 'LaTeX' format")
        if self.fm_type == OMML and \
                (OMML, MathML) in Transforms and (MathML, LaTeX) in Transforms:
            self.formula[MathML] = Transforms[(OMML, MathML)](
                self.formula[OMML], display=self.fm_format
            )
            self.formula[LaTeX] = Transforms[(MathML, LaTeX)](
                self.formula[MathML], display=self.fm_format
            )
            if to_string:
                return self.formula[LaTeX]
            else:
                return None
        if self.fm_type == MathML and \
                (MathML, OMML) in Transforms and (OMML, LaTeX) in Transforms:
            self.formula[OMML] = Transforms[(MathML, OMML)](
                self.formula[MathML], display=self.fm_format
            )
            self.formula[LaTeX] = Transforms[(OMML, LaTeX)](
                self.formula[OMML], display=self.fm_format
            )
            if to_string:
                return self.formula[LaTeX]
            else:
                return None
        raise ValueError("Can't get the formula of 'LaTeX' format")

    def get_mathml(self, to_string=False, safe_mode=False):
        if self.fm_type == MathML:
            if to_string:
                return etree.tostring(self.formula[MathML], encoding=str)
            else:
                return None
        if (self.fm_type, MathML) in Transforms:
            self.formula[MathML] = Transforms[(self.fm_type, MathML)](
                self.formula[self.fm_type], display=self.fm_format
            )
            if to_string:
                return etree.tostring(self.formula[MathML], encoding=str)
            else:
                return None
        if safe_mode:
            raise ValueError("Can't get the formula of 'MathML' format")
        if self.fm_type == LaTeX and \
                (LaTeX, OMML) in Transforms and (OMML, MathML) in Transforms:
            self.formula[OMML] = Transforms[(LaTeX, OMML)](
                self.formula[LaTeX], display=self.fm_format
            )
            self.formula[MathML] = Transforms[(OMML, MathML)](
                self.formula[OMML], display=self.fm_format
            )
            if to_string:
                return etree.tostring(self.formula[MathML], encoding=str)
            else:
                return None
        if self.fm_type == OMML and \
                (OMML, LaTeX) in Transforms and (LaTeX, MathML) in Transforms:
            self.formula[LaTeX] = Transforms[(OMML, LaTeX)](
                self.formula[OMML], display=self.fm_format
            )
            self.formula[MathML] = Transforms[(LaTeX, MathML)](
                self.formula[MathML], display=self.fm_format
            )
            if to_string:
                return etree.tostring(self.formula[MathML], encoding=str)
            else:
                return None
        raise ValueError("Can't get the formula of 'MathML' format")

    def get_omml(self, to_string=False, safe_mode=False):
        if self.fm_type == OMML:
            if to_string:
                return etree.tostring(self.formula[OMML], encoding=str)
            else:
                return None
        if (self.fm_type, OMML) in Transforms:
            self.formula[OMML] = Transforms[(self.fm_type, OMML)](
                self.formula[self.fm_type], display=self.fm_format
            )
            if to_string:
                return etree.tostring(self.formula[OMML], encoding=str)
            else:
                return None
        if safe_mode:
            raise ValueError("Can't get the formula of 'OMML' format")
        if self.fm_type == LaTeX and \
                (LaTeX, MathML) in Transforms and (MathML, OMML) in Transforms:
            self.formula[MathML] = Transforms[(LaTeX, MathML)](
                self.formula[LaTeX], display=self.fm_format
            )
            self.formula[OMML] = Transforms[(MathML, OMML)](
                self.formula[MathML], display=self.fm_format
            )
            if to_string:
                return etree.tostring(self.formula[OMML], encoding=str)
            else:
                return None
        if self.fm_type == MathML and \
                (MathML, LaTeX) in Transforms and (LaTeX, OMML) in Transforms:
            self.formula[LaTeX] = Transforms[(MathML, LaTeX)](
                self.formula[MathML], display=self.fm_format
            )
            self.formula[OMML] = Transforms[(LaTeX, OMML)](
                self.formula[LaTeX], display=self.fm_format
            )
            if to_string:
                return etree.tostring(self.formula[OMML], encoding=str)
            else:
                return None
        raise ValueError("Can't get the formula of 'OMML' format")


if __name__ == '__main__':
    from docx import Document
    doc = Document()
    string = '<>'
    para = doc.add_paragraph() + Formula(string)
    doc.save('demo.docx')
